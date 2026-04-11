from PyQt5.QtWidgets import (
    QMainWindow,
    QWidget,
    QVBoxLayout,
    QHBoxLayout,
    QPushButton,
    QTextEdit,
    QProgressBar,
    QMessageBox,
    QFileDialog,
    QLabel,
    QApplication,
    QButtonGroup,
)
from PyQt5.QtCore import Qt, QThread, pyqtSignal
from PyQt5.QtGui import QImage
from pathlib import Path
import tempfile
import re

from ui.drop_area import DropArea
from ui.result_viewer import ResultViewer
from core.file_validator import FileValidator
from core.recognizer import GLMRecognizer
from core.settings_manager import SettingsManager


class RecognitionThread(QThread):
    finished = pyqtSignal(dict)
    error = pyqtSignal(str)

    def __init__(self, recognizer: GLMRecognizer, file_path: str):
        super().__init__()
        self.recognizer = recognizer
        self.file_path = file_path

    def run(self):
        try:
            result = self.recognizer.recognize(self.file_path)
            self.finished.emit(result)
        except Exception as e:
            self.error.emit(str(e))


class MainWindow(QMainWindow):
    def __init__(self):
        super().__init__()
        self.recognizer = GLMRecognizer()
        self.settings = SettingsManager()
        self._current_result = None
        self._temp_files = []
        self._init_ui()
        self._init_styles()
        self._init_clipboard_listener()
        self._load_settings()

    def _load_settings(self):
        saved_mode = self.settings.get("display_mode", ResultViewer.MODE_MARKDOWN)
        self.result_view.set_mode(saved_mode)
        self._update_mode_buttons(saved_mode)

    def _init_clipboard_listener(self):
        self.setFocusPolicy(Qt.StrongFocus)

    def keyPressEvent(self, event):
        if event.key() == Qt.Key_V and event.modifiers() == Qt.ControlManager:
            self._paste_image()
        else:
            super().keyPressEvent(event)

    def _paste_image(self):
        clipboard = QApplication.clipboard()
        mime_data = clipboard.mimeData()

        if mime_data.hasImage():
            image = clipboard.image()
            if not image.isNull():
                temp_path = tempfile.mktemp(suffix=".png")
                image.save(temp_path)
                self._temp_files.append(temp_path)
                self._update_status("已从剪贴板粘贴图片")
                self._start_recognition(temp_path)
                return True
        elif mime_data.hasUrls():
            urls = mime_data.urls()
            if urls:
                file_path = urls[0].toLocalFile()
                self._update_status(f"已从剪贴板粘贴: {Path(file_path).name}")
                self._start_recognition(file_path)
                return True
        return False

    def _init_ui(self):
        self.setWindowTitle("GLM-OCR 拖拽识别工具")
        self.setMinimumSize(900, 600)

        central = QWidget()
        self.setCentralWidget(central)
        layout = QVBoxLayout(central)

        top_layout = QHBoxLayout()
        self.drop_area = DropArea()
        self.drop_area.file_dropped.connect(self._on_file_dropped)
        self.drop_area.status_changed.connect(self._update_status)
        top_layout.addWidget(self.drop_area, stretch=2)

        btn_layout = QVBoxLayout()
        self.btn_browse = QPushButton("浏览文件")
        self.btn_browse.clicked.connect(self._browse_file)
        self.btn_paste = QPushButton("粘贴图片")
        self.btn_paste.clicked.connect(self._paste_image)
        self.btn_clear = QPushButton("清空")
        self.btn_clear.clicked.connect(self._clear_result)

        for btn in [self.btn_browse, self.btn_paste, self.btn_clear]:
            btn.setFixedHeight(40)
            btn_layout.addWidget(btn)
        btn_layout.addStretch()
        top_layout.addLayout(btn_layout, stretch=1)

        layout.addLayout(top_layout)

        mode_layout = QHBoxLayout()
        mode_label = QLabel("显示模式:")
        mode_label.setStyleSheet("font-weight: bold;")
        mode_layout.addWidget(mode_label)

        self.mode_group = QButtonGroup()
        self.btn_mode_plain = QPushButton("纯文本")
        self.btn_mode_markdown = QPushButton("Markdown")
        self.btn_mode_raw = QPushButton("原始")

        self.mode_group.addButton(self.btn_mode_plain, 0)
        self.mode_group.addButton(self.btn_mode_markdown, 1)
        self.mode_group.addButton(self.btn_mode_raw, 2)

        self.btn_mode_plain.clicked.connect(
            lambda: self._set_display_mode(ResultViewer.MODE_PLAIN)
        )
        self.btn_mode_markdown.clicked.connect(
            lambda: self._set_display_mode(ResultViewer.MODE_MARKDOWN)
        )
        self.btn_mode_raw.clicked.connect(
            lambda: self._set_display_mode(ResultViewer.MODE_RAW)
        )

        for btn in [self.btn_mode_plain, self.btn_mode_markdown, self.btn_mode_raw]:
            btn.setCheckable(True)
            btn.setFixedHeight(30)
            mode_layout.addWidget(btn)

        mode_layout.addStretch()

        self.btn_copy = QPushButton("复制结果")
        self.btn_copy.clicked.connect(self._copy_result)
        self.btn_copy.setFixedHeight(30)
        self.btn_copy.setEnabled(False)
        mode_layout.addWidget(self.btn_copy)

        self.btn_export = QPushButton("导出")
        self.btn_export.clicked.connect(self._export_result)
        self.btn_export.setFixedHeight(30)
        self.btn_export.setEnabled(False)
        mode_layout.addWidget(self.btn_export)

        layout.addLayout(mode_layout)

        self.result_view = ResultViewer()
        layout.addWidget(self.result_view)

        self.status_label = QLabel("就绪")
        self.progress = QProgressBar()
        self.progress.setVisible(False)

        status_layout = QHBoxLayout()
        status_layout.addWidget(self.status_label)
        status_layout.addWidget(self.progress)
        layout.addLayout(status_layout)

    def _update_mode_buttons(self, mode: str):
        if mode == ResultViewer.MODE_PLAIN:
            self.btn_mode_plain.setChecked(True)
        elif mode == ResultViewer.MODE_MARKDOWN:
            self.btn_mode_markdown.setChecked(True)
        else:
            self.btn_mode_raw.setChecked(True)

    def _set_display_mode(self, mode: str):
        self.result_view.set_mode(mode)
        self.settings.set("display_mode", mode)
        self._update_mode_buttons(mode)
        if self._current_result:
            self._show_result(self._current_result)

    def _init_styles(self):
        self.setStyleSheet("""
            QMainWindow { background: #ffffff; }
            QPushButton {
                background: #0078d7;
                color: white;
                border: none;
                border-radius: 4px;
                padding: 8px 16px;
                font-weight: 500;
            }
            QPushButton:hover { background: #005a9e; }
            QPushButton:disabled { background: #ccc; }
            QPushButton:checked {
                background: #005a9e;
                border: 2px solid #003d6b;
            }
            QProgressBar {
                border: 1px solid #ddd;
                border-radius: 3px;
                text-align: center;
            }
            QProgressBar::chunk {
                background: #0078d7;
                width: 10px;
            }
        """)

    def _on_file_dropped(self, file_path: str):
        valid, error = FileValidator.validate(file_path)
        if not valid:
            QMessageBox.warning(self, "格式错误", error)
            return

        self._start_recognition(file_path)

    def _browse_file(self):
        file_path, _ = QFileDialog.getOpenFileName(
            self, "选择文件", "", "图片与PDF (*.jpg *.jpeg *.png *.pdf)"
        )
        if file_path:
            self._on_file_dropped(file_path)

    def _start_recognition(self, file_path: str):
        self.progress.setVisible(True)
        self.progress.setRange(0, 0)
        self.status_label.setText("识别中...")
        self.btn_export.setEnabled(False)
        self.btn_copy.setEnabled(False)

        self.thread = RecognitionThread(self.recognizer, file_path)
        self.thread.finished.connect(self._on_recognition_done)
        self.thread.error.connect(self._on_recognition_error)
        self.thread.start()

    def _on_recognition_done(self, result: dict):
        self.progress.setVisible(False)
        self.status_label.setText("识别完成")
        self._current_result = result
        self._show_result(result)
        self.btn_export.setEnabled(True)
        self.btn_copy.setEnabled(True)

    def _show_result(self, result: dict):
        current_mode = self.result_view._current_mode
        self.result_view.set_result(result["markdown"], mode=current_mode)

    def _on_recognition_error(self, error_msg: str):
        self.progress.setVisible(False)
        self.status_label.setText("识别失败")
        QMessageBox.critical(self, "识别错误", f"识别过程出错:\n{error_msg}")

    def _clear_result(self):
        self.result_view.clear()
        self._current_result = None
        self.status_label.setText("已清空")
        self.btn_export.setEnabled(False)
        self.btn_copy.setEnabled(False)

    def _copy_result(self):
        if not self._current_result:
            return
        text = self.result_view.get_current_text()
        clipboard = QApplication.clipboard()
        clipboard.setText(text)
        self.status_label.setText("已复制到剪贴板")

    def _export_result(self):
        if not self._current_result:
            return

        save_path, _ = QFileDialog.getSaveFileName(
            self,
            "保存结果",
            "glm_ocr_result.md",
            "Markdown文件 (*.md);;文本文件 (*.txt);;JSON文件 (*.json);;HTML文件 (*.html);;Word文档 (*.docx);;CSV文件 (*.csv)",
        )
        if not save_path:
            return

        try:
            ext = Path(save_path).suffix.lower()
            content = self._current_result["markdown"]

            if ext == ".json":
                import json

                with open(save_path, "w", encoding="utf-8") as f:
                    json.dump(self._current_result, f, ensure_ascii=False, indent=2)
            elif ext == ".html":
                html_content = self._generate_html(content)
                with open(save_path, "w", encoding="utf-8") as f:
                    f.write(html_content)
            elif ext == ".docx":
                self._export_docx(content, save_path)
            elif ext == ".csv":
                csv_content = self._generate_csv(content)
                with open(save_path, "w", encoding="utf-8") as f:
                    f.write(csv_content)
            elif ext == ".md":
                with open(save_path, "w", encoding="utf-8") as f:
                    f.write(content)
            else:
                plain_text = self._extract_plain_text(content)
                with open(save_path, "w", encoding="utf-8") as f:
                    f.write(plain_text)

            self.status_label.setText(f"已导出: {Path(save_path).name}")
        except Exception as e:
            QMessageBox.critical(self, "导出失败", str(e))

    def _generate_html(self, markdown: str) -> str:
        from ui.result_viewer import ResultViewer

        viewer = ResultViewer()
        html = viewer._format_markdown(markdown)
        return f"""<!DOCTYPE html>
<html>
<head>
    <meta charset="utf-8">
    <title>GLM-OCR Result</title>
    <style>
        body {{ font-family: Arial, sans-serif; margin: 20px; }}
        pre {{ background: #f4f4f4; padding: 10px; border-radius: 4px; }}
        code {{ background: #f4f4f4; padding: 2px 4px; border-radius: 3px; }}
        h1 {{ color: #333; border-bottom: 1px solid #eee; }}
        h2 {{ color: #555; }}
        h3 {{ color: #666; }}
        li {{ margin-left: 20px; }}
    </style>
</head>
<body>
{html}
</body>
</html>"""

    def _export_docx(self, markdown: str, save_path: str):
        from docx import Document
        from docx.shared import Pt, Inches
        from docx.enum.text import WD_ALIGN_PARAGRAPH

        doc = Document()
        doc.add_heading("GLM-OCR 识别结果", 0)

        lines = markdown.split("\n")
        in_code_block = False
        code_lines = []

        for line in lines:
            if line.startswith("```"):
                if not in_code_block:
                    in_code_block = True
                    code_lines = []
                else:
                    para = doc.add_paragraph()
                    para.add_run("\n".join(code_lines)).font.name = "Courier New"
                    in_code_block = False
                continue

            if in_code_block:
                code_lines.append(line)
                continue

            if line.startswith("# "):
                doc.add_heading(line[2:], 1)
            elif line.startswith("## "):
                doc.add_heading(line[3:], 2)
            elif line.startswith("### "):
                doc.add_heading(line[4:], 3)
            elif line.startswith("- ") or line.startswith("* "):
                doc.add_paragraph(line[2:], style="List Bullet")
            elif re.match(r"^\d+\.\s", line):
                doc.add_paragraph(re.sub(r"^\d+\.\s", "", line), style="List Number")
            elif line.strip():
                doc.add_paragraph(line)

        doc.save(save_path)

    def _generate_csv(self, markdown: str) -> str:
        import csv
        import io

        output = io.StringIO()
        writer = csv.writer(output)
        writer.writerow(["内容"])

        lines = markdown.split("\n")
        for line in lines:
            line = line.strip()
            if line and not line.startswith("#") and not line.startswith("```"):
                writer.writerow([line])

        return output.getvalue()

    def _extract_plain_text(self, text: str) -> str:
        text = re.sub(r"```[\s\S]*?```", "", text)
        text = re.sub(r"`[^`]+`", "", text)
        text = re.sub(r"\*\*([^*]+)\*\*", r"\1", text)
        text = re.sub(r"\*([^*]+)\*", r"\1", text)
        text = re.sub(r"#{1,6}\s+", "", text)
        lines = text.split("\n")
        result = []
        for line in lines:
            line = line.strip()
            if line:
                result.append(line)
        return "\n".join(result)

    def _update_status(self, msg: str):
        self.status_label.setText(msg)
